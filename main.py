import csv, os
from docx import Document

def donnerCours(genie: str, chemin_csv: str = "horaire.csv") -> dict:
    """
    Lit le fichier CSV et retourne les cours du génie spécifié sous forme de dictionnaire.
    """
    horaire_filtré = {}

    with open(chemin_csv, encoding="utf-8") as fichier:
        reader = csv.reader(fichier)
        for row in reader:
            if "ï¿½" in row[-1]:
                row[-1] = row[-1].replace("ï¿½", "a")
            sigle = row[0]
            entry = {
                "Type": row[1],
                "Groupe": row[2],
                "Jour": row[3],
                "Debut": row[4],
                "Duree": row[5],
                "Fin": row[6],
                "Frequence": row[7],
                "Examen": row[8] if row[8] != "Voir note de haut de page" else "Pas de date"
            }
            if sigle.startswith(genie.upper()):
                if sigle not in horaire_filtré:
                    horaire_filtré[sigle] = []
                horaire_filtré[sigle].append(entry)
    
    return horaire_filtré


def generer_docx(horaire: dict, titre: str = "Horaire") -> None:
    """
    Crée un fichier Word contenant l'horaire sous forme de tableau.
    """
    document = Document()
    document.add_heading(titre, 0)

    for sigle, cours_list in horaire.items():
        document.add_heading(sigle, level=1)

        # Créer une table avec les entêtes
        table = document.add_table(rows=1, cols=8)
        cells = table.rows[0].cells
        cells[0].text = "Type"
        cells[1].text = "Groupe"
        cells[2].text = "Jour"
        cells[3].text = "Début"
        cells[4].text = "Durée"
        cells[5].text = "Fin"
        cells[6].text = "Fréquence"
        cells[7].text = "Examen"

        for cours in cours_list:
            row_cells = table.add_row().cells
            row_cells[0].text = cours["Type"]
            row_cells[1].text = cours["Groupe"]
            row_cells[2].text = cours["Jour"]
            row_cells[3].text = cours["Debut"]
            row_cells[4].text = cours["Duree"]
            row_cells[5].text = cours["Fin"]
            row_cells[6].text = cours["Frequence"]
            row_cells[7].text = cours["Examen"]

    document.save(f"{titre}.docx")

def ouvrirDoc(chemin: str):
    if os.path.exists(chemin):
        os.startfile(chemin)
    else:
        print(f"Le fichier '{chemin}' n'existe pas.")

def fusionner_docs(fichier1: str, fichier2: str, fichier_sortie: str):
    doc1 = Document(fichier1)
    doc2 = Document(fichier2)

    doc1.add_page_break()

    for element in doc2.element.body:
        doc1.element.body.append(element)

    doc1.save(fichier_sortie)

def fusionner_selection_dossiers(dossier: str = ".", sortie: str = "Fusion_result.docx"):
    fichiers_docx = [f for f in os.listdir(dossier) if f.endswith(".docx")]
    
    if not fichiers_docx:
        print("Aucun fichier .docx trouvé dans le dossier.")
        return

    print("Fichiers disponibles :")
    for i, fichier in enumerate(fichiers_docx):
        print(f"{i + 1}. {fichier}")

    print("\nEntrez les numéros des fichiers à fusionner, séparés par des virgules (ex: 1,3,4) :")
    while True:
        try:
            choix = input("> ")
            indices = [int(i.strip()) - 1 for i in choix.split(",")]
            fichiers_choisis = [fichiers_docx[i] for i in indices]
            if len(fichiers_choisis) < 2:
                raise ValueError("Il faut au moins deux fichiers pour fusionner.")
            break
        except (ValueError, IndexError):
            print("Entrée invalide. Réessaie avec au moins deux numéros valides.")

    try:
        doc_fusion = Document(os.path.join(dossier, fichiers_choisis[0]))

        for fichier in fichiers_choisis[1:]:
            doc_fusion.add_page_break()
            doc_ajout = Document(os.path.join(dossier, fichier))
            for element in doc_ajout.element.body:
                doc_fusion.element.body.append(element)

        chemin_sortie = os.path.join(dossier, sortie)
        doc_fusion.save(chemin_sortie)
        print(f"\n✅ Fichiers fusionnés dans : {chemin_sortie}")
        os.startfile(chemin_sortie)

    except Exception as e:
        print(f"Erreur pendant la fusion : {e}")


def run():
    print(f"Bienvenue dans le générateur de tableau d'horaire!\nVoici les commandes disponibles: ")
    while True:
        try:
            n=int(input(f"1. Donner un document contenant les cours d'un genie spécifique.\n2. Fusionner deux documents en un seul. \n3. Effacer un document\n4. Quitter\n"))
            if n>4 or n<1:
                raise ValueError("Le nombre doit être entre 1 et 4 inclut!\n")
        except ValueError as e:
            print(f"Erreur: {e}")
        
        match n:
            case 1:
                cours = input("Donne le cours que tu veux: ")
                titre = input("Donne un titre à ton document: ")
                horaire = donnerCours(cours)
                generer_docx(horaire, titre)

                while True:
                    yes_or_no = input("Tu veux que je t'ouvre le document ? (yes/no): ").lower()
                    if yes_or_no in ["yes", "no"]:
                        break
                    else:
                        print("Erreur : Réponds uniquement par 'yes' ou 'no'.")
                
                if yes_or_no == "yes":
                    ouvrirDoc(f"{titre}.docx")
            case 2:
                try:
                    fusionner_selection_dossiers(".")  # ou "." pour le dossier courant
                except Exception as e:
                    print(f"Erreur: {e}")

            case 3:
                try:
                    doc = input("Donne le nom du document à enlever : ")
                    if not doc.endswith(".docx"):
                        doc += ".docx"

                    if not os.path.exists(doc):
                        raise FileNotFoundError("Le document n'existe pas.")

                    os.remove(doc)
                    print("✅ Document supprimé avec succès.")
                except FileNotFoundError as e:
                    print(f"Erreur : {e}")
                except Exception as e:
                    print(f"Une erreur est survenue : {e}")


            case 4:
                print("Au revoir!")
                break
        

if __name__== "__main__":
    run()